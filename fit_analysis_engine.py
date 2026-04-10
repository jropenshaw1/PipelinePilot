# fit_analysis_engine.py — PipelinePilot fit analysis engine
#
# Mirrors the Streamlit fit_analysis.py approach:
#   - Single Claude API call returns all artifacts in one JSON
#   - Tailored resume = master resume copy + tailored summary prepended
#   - Same document styling as the Streamlit output
#
# Runs in a background thread. Never blocks the UI.
# Writes to opportunity folder:
#   - fit_analysis.md              (YAML → SQLite index)
#   - FitAnalysis_{Company}.docx
#   - CoverLetter_{Company}_{Role}.docx
#   - Resume_{Company}_{Role}.docx
#   - InterviewGuide_{Company}_{Role}.docx

import copy
import json
import os
import re
import shutil
import threading
import yaml
from datetime import date, datetime
from pathlib import Path

import database
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# ─────────────────────────────────────────────
# Filename sanitization
# Mirrors filesystem.py sanitize_component():
# strips spaces, commas, periods, slashes,
# ampersands, and any non-alphanumeric chars.
# ─────────────────────────────────────────────

_STRIP_CHARS = re.compile(r"[&,./\\]")
_KEEP_CHARS  = re.compile(r"[^A-Za-z0-9\-]")


def _sanitize(text: str) -> str:
    text = text.strip().replace(" ", "")
    text = _STRIP_CHARS.sub("", text)
    text = _KEEP_CHARS.sub("", text)
    return text


# ─────────────────────────────────────────────
# Config
# ─────────────────────────────────────────────

MODEL = "claude-sonnet-4-20250514"
MAX_TOKENS = 8000

# Stage definitions for the progress dialog.
# Each tuple: (stage_id, display_label, subtitle)
ANALYSIS_STAGES = [
    ("read_inputs",     "Read Inputs",      "JD + Resume"),
    ("api_call",        "AI Analysis",      "All agents (single call)"),
    ("fit_summary",     "Fit Summary",      "Score + alignment report"),
    ("cover_letter",    "Cover Letter",     "Personalized letter"),
    ("resume",          "Tailored Resume",  "ATS-ready copy"),
    ("interview_guide", "Interview Guide",  "Prep questions + strategies"),
]

# Default API key location — same as Streamlit version
_API_KEY_FILENAME = "ClaudeAPIkey.txt"


def _get_api_key(job_search_root: str, override_key: str = "") -> str:
    """Return API key from override, key file, or raise."""
    if override_key and override_key.strip():
        return override_key.strip()
    key_file = Path(job_search_root) / _API_KEY_FILENAME
    if key_file.exists():
        return key_file.read_text(encoding="utf-8").strip()
    raise RuntimeError(
        f"Claude API key not found.\n\n"
        f"Expected at: {key_file}\n\n"
        "Either add ClaudeAPIkey.txt to your job search root folder, "
        "or set the key in Settings → Anthropic API Key."
    )


# ─────────────────────────────────────────────
# Personal context skill loader
# Reads a SKILL.md from disk and strips YAML
# frontmatter, returning only the body text.
# Returns empty string silently if path is
# not configured, missing, or unreadable.
# ─────────────────────────────────────────────

def _load_context_skill(skill_path: str) -> str:
    """Load personal context from a SKILL.md file if path is configured."""
    if not skill_path or not skill_path.strip():
        return ""
    path = Path(skill_path.strip())
    if not path.exists():
        return ""
    try:
        content = path.read_text(encoding="utf-8")
        # Strip YAML frontmatter block (--- ... ---)
        if content.startswith("---"):
            parts = content.split("---", 2)
            if len(parts) >= 3:
                content = parts[2].strip()
        return content
    except Exception:
        return ""


# ─────────────────────────────────────────────
# Prompts — single call returns full JSON
# ─────────────────────────────────────────────

SYSTEM_PROMPT = """You are a senior career fit analyst using a dual-voice framework.

You hold two perspectives simultaneously:
- Advocate: Finds the best honest case for the candidacy. Highlights transferable skills, reframes experience positively.
- Auditor: Tells the hard truth. Where gaps are real, where the candidate would be stretching credibility.

Scoring scale:
0.0–0.3: Significant gaps. Long shot.
0.4–0.6: Partial fit. Gaps exist but may be addressable.
0.7–0.8: Strong fit. Minor gaps or growth areas.
0.9–1.0: Near-perfect alignment. (Be suspicious — recheck.)

Never inflate a score to be encouraging. Use 0.05 increments only (0.55, 0.70, 0.80 — never 0.67 or 0.83).

IMPORTANT — Degree requirements: Many JDs state "Bachelor's degree OR equivalent combination of education and experience."
This is NOT a hard degree requirement.
Treat "equivalent combination" language as satisfied by extensive relevant experience.
Only flag degree_required as a genuine gap if the JD explicitly states the degree is mandatory with no equivalency option.

Interview guide: Generate 8–10 targeted interview questions the candidate is likely to face,
each with a recommended answer strategy that draws on their specific experience.
Mix behavioral (STAR format), technical, and situational questions relevant to the role.

Cover letter: Write a full, specific, professional cover letter (3-4 paragraphs). No placeholders.
Reference actual accomplishments from the resume.

HARD VOICE BANS — these are enforcement rules, not style suggestions. A violation means rewrite the sentence before returning the JSON.

1. No generic openers. Banned phrases: "I am writing to express my interest in," "I am excited to apply for," "I am pleased to submit my application for," and any variant that announces the act of applying. Open with substance — a specific accomplishment or a concrete claim the candidate brings. The opening sentence must contain a fact about the candidate, not a fact about the application.

2. No superlative matching. Do NOT write any sentence whose purpose is to claim alignment between the candidate and the role, company, or stated needs. Banned patterns include: "aligns with," "aligns directly with," "aligns perfectly with," "represents exactly," "is an ideal match for," "uniquely positioned to," "precisely the kind of," "is a natural fit for," "essential for," "required to," "needed for," "proven track record of," and all synonym variations. The rule is structural: any sentence that tells the reader the candidate matches, fits, maps to, demonstrates, or proves readiness for the role is a violation. State the accomplishment. Stop. Let the reader draw the connection.

3. Stay in the candidate's lane. Write what the candidate brings, not what the company needs. Do NOT speculate about what the company "requires," "needs," "is looking for," or "seeks." Do NOT invent a gap and frame the candidate as filling it. Do NOT write "[Company] seeks..." or "[Company]'s need for..." or "to address your..." — the hiring team wrote the JD; they know their needs. If a documented JD requirement is directly relevant, reference it once with tentative framing ("the JD notes...," "the role calls for...") rather than confident framing. Default to silence on company needs.

4. No generic closers. Banned: "contribute to [Company]'s continued growth," "contribute to your operational excellence," "help [Company] achieve its goals," "be a valuable addition to your team," and any variant that promises vague future contribution. Close with something specific the candidate wants to discuss or a direct statement of interest — not a platitude about the company's future.

5. No em dashes anywhere in the cover letter. Use commas, colons, or new sentences.

CRITICAL — Job title accuracy: The authoritative job title is provided at the top of the user message as COMPANY and ROLE. Use the ROLE value verbatim when referencing the position in the cover letter. Do not paraphrase, shorten, reword, or infer a different title from the JD text.

Self-check before returning JSON: Scan the generated cover_letter field for the five ban categories above. If any banned phrase, pattern, or em dash appears, rewrite the offending sentence before returning. A banned phrase that survives to the response is a failure of this task.

Candidate contact: If the PERSONAL CONTEXT section below contains contact details (full name,
email, phone, location, LinkedIn), extract them and populate the candidate_contact object in
the JSON response. The personal context is the authoritative source for these values — do NOT
parse them from the resume. If personal context is absent or a specific field is missing,
return an empty string for that field.

You must respond with a JSON object and nothing else. No markdown, no preamble.
The JSON must have exactly this structure:
{
  "role_summary": "2-3 sentence summary of the role",
  "fit_score": 0.00,
  "score_justification": "One sentence explaining the score",
  "alignment": [
    {"area": "area name", "detail": "how experience maps"}
  ],
  "gaps": [
    {"area": "gap area", "severity": "minor|moderate|significant", "detail": "specific gap description"}
  ],
  "advocate_case": "Best honest argument for applying",
  "auditor_case": "Honest risks and concerns",
  "next_steps": ["step 1", "step 2", "step 3"],
  "cover_letter": "Full cover letter text, professional, 3-4 paragraphs, no placeholders",
  "candidate_contact": {
    "name": "Full name from personal context, or empty string",
    "email": "Email from personal context, or empty string",
    "phone": "Phone from personal context, or empty string",
    "location": "City and state from personal context, or empty string",
    "linkedin": "LinkedIn URL from personal context, or empty string"
  },
  "resume_highlights": ["bullet 1 tailored to this JD", "bullet 2", "bullet 3", "bullet 4", "bullet 5"],
  "interview_guide": [
    {"question": "question text", "type": "behavioral|technical|situational", "strategy": "recommended answer approach drawing on candidate's specific experience"}
  ]
}"""


# ─────────────────────────────────────────────
# Text extraction
# ─────────────────────────────────────────────

def _read_text_file(path: Path) -> str:
    """Read a text file, falling back to latin-1 if UTF-8 decoding fails.
    Handles Windows-1252 characters common in LinkedIn/browser copy-paste.
    """
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="latin-1")


def extract_docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts)


def find_jd_file(folder_path: Path) -> Path:
    name = folder_path.name
    skip = ("resume", "cover", "interview", "fitanalysis", "interviewguide")

    candidates = [
        folder_path / f"JD_{name}.docx",
        *folder_path.glob("JD_*.docx"),
        *[f for f in folder_path.glob("*.docx")
          if not any(s in f.name.lower().replace("_", "") for s in skip)],
        *folder_path.glob("*.txt"),
    ]
    for c in candidates:
        if c.exists():
            return c

    raise FileNotFoundError(
        f"No job description file found in {name}.\n\n"
        "Paste the job description into the JD document in the opportunity folder, then try again."
    )


# ─────────────────────────────────────────────
# API call
# company and role are passed explicitly so the model has the authoritative
# values as named facts — not something it must infer from JD text.
# ─────────────────────────────────────────────

def _call_api(client, jd_text: str, resume_text: str, company: str, role: str,
              personal_context: str = "") -> dict:
    effective_prompt = SYSTEM_PROMPT
    if personal_context:
        effective_prompt = (
            SYSTEM_PROMPT
            + "\n\n---\n\n# PERSONAL CONTEXT\n\n"
            + personal_context
            + "\n\n---"
        )
    message = client.messages.create(
        model=MODEL,
        max_tokens=MAX_TOKENS,
        temperature=0,
        system=effective_prompt,
        messages=[{
            "role": "user",
            "content": (
                f"COMPANY: {company}\n"
                f"ROLE: {role}\n\n"
                f"RESUME:\n{resume_text}\n\n---\n\nJOB DESCRIPTION:\n{jd_text}"
            )
        }]
    )
    raw = message.content[0].text.strip()
    raw = re.sub(r"^```json\s*", "", raw)
    raw = re.sub(r"```$", "", raw.strip())
    try:
        result = json.loads(raw.strip())
    except json.JSONDecodeError as e:
        raise RuntimeError(f"API returned invalid JSON: {e}\n\nFirst 500 chars:\n{raw[:500]}")

    # Snap fit_score to 0.05 increments
    raw_score = float(result.get("fit_score", 0.5))
    result["fit_score"] = round(round(raw_score / 0.05) * 0.05, 2)

    return result


# ─────────────────────────────────────────────
# Document helpers (match Streamlit styling)
# ─────────────────────────────────────────────

def _heading(doc, text, level=1, color=None):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def _body(doc, text):
    p = doc.add_paragraph(text)
    p.style = "Normal"
    return p


def _bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


# ─────────────────────────────────────────────
# Document generators
# ─────────────────────────────────────────────

def generate_fit_summary(result: dict, company: str, role: str, folder_path: Path) -> Path:
    doc = Document()
    doc.core_properties.author = ""
    now_str = datetime.now().strftime("%B %d, %Y")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"Fit Analysis — {company}")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1F, 0x7A, 0x8C)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"{role}  ·  {now_str}")
    doc.add_paragraph()

    score = result["fit_score"]
    score_color = (0x1A, 0x7A, 0x4A) if score >= 0.6 else (0xC0, 0x39, 0x2B)
    p = doc.add_paragraph()
    run = p.add_run(f"Fit Score: {score:.2f}  —  {result['score_justification']}")
    run.bold = True
    run.font.color.rgb = RGBColor(*score_color)

    _heading(doc, "Role Summary", 2)
    _body(doc, result["role_summary"])

    _heading(doc, "Alignment", 2)
    for a in result.get("alignment", []):
        p = doc.add_paragraph()
        p.add_run(f"{a['area']}: ").bold = True
        p.add_run(a["detail"])

    _heading(doc, "Gaps", 2)
    severity_colors = {
        "significant": (0xC0, 0x39, 0x2B),
        "moderate":    (0xC0, 0x7A, 0x00),
        "minor":       (0x55, 0x55, 0x55),
    }
    for g in result.get("gaps", []):
        p = doc.add_paragraph()
        run = p.add_run(f"[{g['severity'].upper()}] {g['area']}: ")
        run.bold = True
        run.font.color.rgb = RGBColor(*severity_colors.get(g["severity"], (0x55, 0x55, 0x55)))
        p.add_run(g["detail"])

    _heading(doc, "Advocate Case", 2)
    _body(doc, result["advocate_case"])

    _heading(doc, "Auditor Case", 2)
    _body(doc, result["auditor_case"])

    _heading(doc, "Next Steps", 2)
    for step in result.get("next_steps", []):
        _bullet(doc, step)

    doc.add_paragraph()
    _heading(doc, "⚠️ Before You Apply", 2)
    _bullet(doc, "Cover letter — verify the contact header and signature name are accurate before sending.")
    _bullet(doc, "Review all AI-generated content — read the cover letter and tailored resume carefully before submitting. Do not send unreviewed AI output to a potential employer.")
    _bullet(doc, "If you update your resume based on any gap items above, rerun this analysis before applying.")

    out = folder_path / f"FitAnalysis_{_sanitize(company)}.docx"
    doc.save(str(out))
    return out


def generate_cover_letter(result: dict, company: str, role: str, folder_path: Path) -> Path:
    doc = Document()
    doc.core_properties.author = ""

    # Tighten margins to maximize content area for one-page fit.
    for section in doc.sections:
        section.top_margin    = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin   = Inches(0.9)
        section.right_margin  = Inches(0.9)

    now_str = datetime.now().strftime("%B %d, %Y")

    # Contact header — populated from personal context skill if available.
    # If candidate_contact is missing or empty, header is omitted (legacy behavior).
    contact = result.get("candidate_contact") or {}
    contact_name     = (contact.get("name")     or "").strip()
    contact_email    = (contact.get("email")    or "").strip()
    contact_phone    = (contact.get("phone")    or "").strip()
    contact_location = (contact.get("location") or "").strip()
    contact_linkedin = (contact.get("linkedin") or "").strip()

    if contact_name:
        name_p = doc.add_paragraph()
        name_p.paragraph_format.space_after = Pt(2)
        name_run = name_p.add_run(contact_name)
        name_run.bold = True
        name_run.font.size = Pt(13)

        # Single-line contact strip with bullet separators (no em dashes)
        contact_parts = [p for p in (contact_location, contact_phone, contact_email) if p]
        if contact_parts:
            cp = doc.add_paragraph(" \u2022 ".join(contact_parts))
            cp.paragraph_format.space_after = Pt(2)
        if contact_linkedin:
            lp = doc.add_paragraph(contact_linkedin)
            lp.paragraph_format.space_after = Pt(10)

    date_p = doc.add_paragraph(now_str)
    date_p.paragraph_format.space_after = Pt(10)

    salutation_p = doc.add_paragraph(f"Dear {company} Hiring Team,")
    salutation_p.paragraph_format.space_after = Pt(6)

    # Strip salutation and closing — we add those structurally.
    # Break at first closing phrase so the signature name after it is also excluded.
    _closing_patterns = ("sincerely", "best regards", "regards,", "warm regards",
                         "kind regards", "respectfully", "yours truly", "best,")
    body_lines = []
    for line in result["cover_letter"].split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.lower().startswith("dear "):
            continue  # Skip salutation — we add our own
        if stripped.lower().startswith(_closing_patterns):
            break  # Closing found — discard everything from here on (including signature)
        body_lines.append(stripped)

    # Render body paragraphs with controlled spacing — no blank paragraphs between.
    for para_text in body_lines:
        body_p = doc.add_paragraph(para_text.strip())
        body_p.style = "Normal"
        body_p.paragraph_format.space_after = Pt(6)

    closing_p = doc.add_paragraph("Sincerely,")
    closing_p.paragraph_format.space_after = Pt(2)
    if contact_name:
        sig_p = doc.add_paragraph(contact_name)
    else:
        sig_p = doc.add_paragraph()  # blank line for signature when no name available
    sig_p.paragraph_format.space_after = Pt(0)

    out = folder_path / f"CoverLetter_{_sanitize(company)}_{_sanitize(role)}.docx"
    doc.save(str(out))
    return out


def generate_tailored_resume(result: dict, company: str, role: str,
                              folder_path: Path, resume_path: str) -> Path:
    """
    Copy master resume as-is. The master resume (V18+) already contains a
    tailored professional summary — no prepend block needed or wanted.
    The resume_highlights field from the API result is surfaced in the
    FitAnalysis doc instead, keeping the resume clean.
    """
    # Output is always .docx for ATS compatibility.
    # If the configured resume_path is .md (API input master), resolve the
    # .docx sibling at the same location for the copy source.
    src = Path(resume_path)
    if src.suffix.lower() != ".docx":
        docx_sibling = src.with_suffix(".docx")
        if not docx_sibling.exists():
            raise FileNotFoundError(
                f"ATS-ready .docx resume not found at {docx_sibling}.\n\n"
                "Place a formatted .docx version alongside your .md resume, "
                "or set Resume Path directly to the .docx file."
            )
        src = docx_sibling
    out = folder_path / f"Resume_{_sanitize(company)}_{_sanitize(role)}.docx"
    shutil.copy2(str(src), str(out))
    return out


def generate_interview_guide(result: dict, company: str, role: str, folder_path: Path) -> Path:
    doc = Document()
    doc.core_properties.author = ""
    now_str = datetime.now().strftime("%B %d, %Y")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"Interview Prep Guide — {company}")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1F, 0x7A, 0x8C)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"{role}  ·  {now_str}")
    doc.add_paragraph()

    type_labels = {
        "behavioral":  ("BEHAVIORAL",  RGBColor(0x1A, 0x6F, 0xA0)),
        "technical":   ("TECHNICAL",   RGBColor(0x1F, 0x7A, 0x8C)),
        "situational": ("SITUATIONAL", RGBColor(0x7A, 0x4A, 0x1A)),
    }

    for i, item in enumerate(result.get("interview_guide", []), 1):
        q_type = item.get("type", "behavioral").lower()
        label, color = type_labels.get(q_type, ("GENERAL", RGBColor(0x55, 0x55, 0x55)))

        p = doc.add_paragraph()
        tag = p.add_run(f"[{label}]  ")
        tag.bold = True
        tag.font.color.rgb = color
        q_run = p.add_run(f"Q{i}: {item['question']}")
        q_run.bold = True

        strat_p = doc.add_paragraph()
        strat_p.paragraph_format.left_indent = Inches(0.25)
        strat_p.add_run("Strategy: ").bold = True
        strat_p.add_run(item.get("strategy", ""))
        doc.add_paragraph()

    out = folder_path / f"InterviewGuide_{_sanitize(company)}_{_sanitize(role)}.docx"
    doc.save(str(out))
    return out


# ─────────────────────────────────────────────
# fit_analysis.md writer
# ─────────────────────────────────────────────

_MD = """\
---
{yaml_block}
---

# Fit Analysis — {company} | {role}

*Generated: {generated_date} | Model: {model}*

---

## Role Summary

{role_summary}

---

## Alignment

{alignment}

---

## Gaps

{gaps}

---

## Advocate Case

{advocate_case}

---

## Auditor Case

{auditor_case}

---

## Fit Score

**{fit_pct}** — {score_justification}

---

## Next Steps

{next_steps}

---

### ⚠️ Before You Apply

- **Cover letter** — verify the contact header and signature name are accurate before sending.
- **Review all AI-generated content** — read the cover letter and tailored resume carefully before submitting. Do not send unreviewed AI output to a potential employer.
- **If you update your resume** based on any gap items above, rerun this analysis before applying.
"""


def _bl(items) -> str:
    return "\n".join(f"- {i}" for i in items) if items else "_(none)_"


def write_fit_analysis_md(folder_path: Path, company: str, role: str,
                           job_url, fit_threshold: float, result: dict) -> Path:
    top_gaps = [g["area"] for g in result.get("gaps", [])[:3]]
    top_strengths = []
    adv = result.get("advocate_case", "")
    if adv:
        top_strengths = [s.strip() for s in adv.split(".") if len(s.strip()) > 20][:3]

    yaml_data = {
        "company": company,
        "role": role,
        "job_url": job_url,
        "fit_score": result["fit_score"],
        "fit_threshold": fit_threshold,
        "recommendation": "APPLY" if result["fit_score"] >= 0.75 else (
            "HOLD" if result["fit_score"] >= 0.60 else "PASS"
        ),
        "top_strengths": top_strengths,
        "top_gaps": top_gaps,
        "generated_date": date.today().isoformat(),
        "model": MODEL,
    }

    alignment_lines = [
        f"- **{a['area']}**: {a['detail']}" for a in result.get("alignment", [])
    ]
    gap_lines = [
        f"- **[{g['severity'].upper()}] {g['area']}**: {g['detail']}"
        for g in result.get("gaps", [])
    ]

    content = _MD.format(
        yaml_block=yaml.dump(yaml_data, default_flow_style=False, allow_unicode=True).strip(),
        company=company, role=role,
        generated_date=date.today().isoformat(), model=MODEL,
        role_summary=result.get("role_summary", ""),
        alignment="\n".join(alignment_lines) if alignment_lines else "_(none)_",
        gaps="\n".join(gap_lines) if gap_lines else "_(none)_",
        advocate_case=result.get("advocate_case", ""),
        auditor_case=result.get("auditor_case", ""),
        fit_pct=f"{result['fit_score']:.0%}",
        score_justification=result.get("score_justification", ""),
        next_steps=_bl(result.get("next_steps", [])),
    )

    out = folder_path / "fit_analysis.md"
    out.write_text(content, encoding="utf-8")
    return out


# ─────────────────────────────────────────────
# Main entry point
# ─────────────────────────────────────────────

def run_fit_analysis(
    db_path: Path,
    job_search_root: str,
    folder_name: str,
    resume_path: str,
    api_key: str,           # empty string = fall back to ClaudeAPIkey.txt
    fit_threshold: float,
    on_success,             # callback(result: dict, output_files: list[Path])
    on_error,               # callback(message: str)
    progress_cb=None,       # callback(step: str, detail: str)
    context_skill_path: str = "",  # path to personal SKILL.md; omit for neutral output
) -> threading.Thread:

    def _p(step, detail="", stage=None):
        if progress_cb:
            progress_cb(step, detail, stage)

    def _worker():
        try:
            import anthropic
        except ImportError:
            on_error("anthropic library not installed.\n\nRun:  pip install anthropic")
            return

        try:
            resolved_key = _get_api_key(job_search_root, api_key)
            client = anthropic.Anthropic(api_key=resolved_key)

            folder_path = Path(job_search_root) / folder_name
            opp = database.get_opportunity(db_path, folder_name)
            if not opp:
                on_error(f"Opportunity '{folder_name}' not found.")
                return

            company = opp["company_name"]
            role = opp["role_title"]
            job_url = opp.get("job_url")

            _p("Reading job description...", stage="read_inputs")
            jd_file = find_jd_file(folder_path)
            jd_text = extract_docx_text(jd_file) if jd_file.suffix == ".docx" \
                else _read_text_file(jd_file)

            if len(jd_text.strip()) < 100:
                on_error(
                    f"Job description appears empty ({len(jd_text.strip())} chars).\n\n"
                    f"Paste the full JD into {jd_file.name} and try again."
                )
                return

            _p("Reading master resume...", stage="read_inputs")
            resume_file = Path(resume_path)
            if not resume_file.exists():
                on_error(f"Resume not found:\n{resume_path}\n\nUpdate Settings → Resume Path.")
                return
            resume_text = extract_docx_text(resume_file) if resume_file.suffix == ".docx" \
                else _read_text_file(resume_file)

            # Single API call — all artifacts in one response.
            # company and role passed explicitly so the model has authoritative
            # values as named facts at the top of the user message.
            _p("Running fit analysis...", "All agents (single call)", stage="api_call")
            personal_context = _load_context_skill(context_skill_path)
            result = _call_api(client, jd_text[:8000], resume_text[:6000], company, role,
                               personal_context=personal_context)

            recommendation = (
                "APPLY" if result["fit_score"] >= 0.75 else
                "HOLD" if result["fit_score"] >= 0.60 else "PASS"
            )

            _p("Writing fit_analysis.md...", stage="fit_summary")
            md_path = write_fit_analysis_md(
                folder_path, company, role, job_url, fit_threshold, result
            )

            _p("Generating fit summary...", stage="fit_summary")
            fa_path = generate_fit_summary(result, company, role, folder_path)

            _p("Generating cover letter...", stage="cover_letter")
            cl_path = generate_cover_letter(result, company, role, folder_path)

            _p("Generating tailored resume...", stage="resume")
            res_path = generate_tailored_resume(
                result, company, role, folder_path, resume_path
            )

            _p("Generating interview prep guide...", stage="interview_guide")
            ip_path = generate_interview_guide(result, company, role, folder_path)

            # Update SQLite
            database.update_opportunity(db_path, folder_name, {
                "fit_score": result["fit_score"],
                "fit_threshold": fit_threshold,
                "recommendation": recommendation,
                "top_strengths": str([
                    s.strip() for s in result.get("advocate_case", "").split(".")
                    if len(s.strip()) > 20
                ][:3]),
                "top_gaps": str([g["area"] for g in result.get("gaps", [])[:3]]),
                "status": "Analyzing",
            })

            output_files = [md_path, fa_path, cl_path, res_path, ip_path]
            _p("Complete.", f"{result['fit_score']:.0%} — {recommendation}", stage="complete")
            on_success(result, output_files)

        except Exception as e:
            try:
                import anthropic
                if isinstance(e, anthropic.AuthenticationError):
                    on_error("Invalid API key.\n\nCheck ClaudeAPIkey.txt or Settings → Anthropic API Key.")
                    return
                if isinstance(e, anthropic.RateLimitError):
                    on_error("Rate limit reached. Wait a moment and try again.")
                    return
            except ImportError:
                pass
            if isinstance(e, (FileNotFoundError, RuntimeError)):
                on_error(str(e))
            else:
                on_error(f"Unexpected error — {type(e).__name__}: {e}")

    t = threading.Thread(target=_worker, daemon=True)
    t.start()
    return t
